#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成生物标志物模块的汇报材料（Word + Excel）。

从本模块既有的只读数据源装配出两份成品文档，便于汇报展示：
- out/report/biomarkers_report.docx  图文说明（模块概述、各项标志物、参考范围与文献、数据局限）
- out/report/biomarkers_report.xlsx  多 sheet 表格（标志物清单、参考范围、S1–S15 队列、文献、局限）

数据来源（均只读，不修改）：
- README.md ......................... 标志物含义、适用范围、数据局限文字
- out/biomarker_reference_ranges.json  参考范围 / reference_type / confidence / 文献
- out/csv/cohort_biomarkers.csv ..... S1–S15 队列每人一行的标签 + 各项数值
- biomarkers.py (BIOMARKER_NAMES) ... 标志物的权威顺序与命名（条数以此为准）
- reasonableness.py (SPECS) ......... 标志物↔标签预期方向、部位、rationale、DISCLAIMER

用法：
    python analysis/02_biomarkers/make_report_docs.py
依赖：python-docx、pandas、openpyxl（写 xlsx 必需，缺则 pip install openpyxl）。
"""
from __future__ import annotations

import json
import os
import sys
from datetime import date

import pandas as pd

# 让脚本既能直接运行又能在包内导入
_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)

from biomarkers import BIOMARKER_NAMES          # noqa: E402  权威标志物顺序
from reasonableness import SPEC_BY_NAME, DISCLAIMER  # noqa: E402

# ----------------------------------------------------------------------------
# 路径
# ----------------------------------------------------------------------------
OUT_DIR = os.path.join(_HERE, "out")
REPORT_DIR = os.path.join(OUT_DIR, "report")
REF_JSON = os.path.join(OUT_DIR, "biomarker_reference_ranges.json")
COHORT_CSV = os.path.join(OUT_DIR, "csv", "cohort_biomarkers.csv")
DOCX_PATH = os.path.join(REPORT_DIR, "biomarkers_report.docx")
XLSX_PATH = os.path.join(REPORT_DIR, "biomarkers_report.xlsx")

# ----------------------------------------------------------------------------
# 模态分组 + 中文含义（取自 README 的三张表，逐项对应）
# ----------------------------------------------------------------------------
MODALITY = {
    # EEG（皮层）
    "pathological_asymmetry_index": "EEG",
    "corticomuscular_coherence_beta": "EEG",
    "prefrontal_theta_beta_ratio": "EEG",
    "interhemispheric_motor_coherence": "EEG",
    "movement_mu_power_change": "EEG",
    "movement_beta_power_change": "EEG",
    # EMG（外周/肌肉）
    "resting_emg_level": "EMG",
    "wrist_co_contraction_index": "EMG",
    "finger_co_contraction_index": "EMG",
    "emg_activation_rms": "EMG",
    "fcr_iemg": "EMG",
    "fds_iemg": "EMG",
    "ecu_iemg": "EMG",
    "extensor_digitorum_iemg": "EMG",
    "flexor_extensor_iemg_ratio": "EMG",
    "emg_burst_duration": "EMG",
    "fcr_mdf": "EMG",
    "fds_mdf": "EMG",
    "ecu_mdf": "EMG",
    "extensor_digitorum_mdf": "EMG",
    # IMU（运动学）
    "movement_smoothness_sparc": "IMU",
    "range_of_motion_proxy": "IMU",
    "tremor_index_3_6hz": "IMU",
    "wrist_flexion_peak_velocity": "IMU",
    "wrist_extension_peak_velocity": "IMU",
    "finger_extension_peak_velocity": "IMU",
}

MEANING_ZH = {
    "pathological_asymmetry_index": "受损 vs 健侧运动皮层 μ/β 功率不对称（PAI）",
    "corticomuscular_coherence_beta": "受损半球–患手主动肌 β 带相干（CMC）",
    "prefrontal_theta_beta_ratio": "前额叶 (Fp1/Fp2/Fz) θ(4–8)/β(13–30) 功率比",
    "interhemispheric_motor_coherence": "健-患侧运动皮层 (C3 簇↔C4 簇) β 带相干",
    "movement_mu_power_change": "运动相关 μ(8–12) 功率变化（去同步为负）",
    "movement_beta_power_change": "运动相关 β(13–30) 功率变化（反弹为正）",
    "resting_emg_level": "患手屈肌静息肌电 RMS（肌张力代理）",
    "wrist_co_contraction_index": "腕屈肌 (FCR)/腕伸肌 (ECU) 共收缩包络重叠（CCI-腕）",
    "finger_co_contraction_index": "指浅屈肌 (FDS)/指伸肌共收缩包络重叠（CCI-指）",
    "emg_activation_rms": "全段自主激活幅度 RMS",
    "fcr_iemg": "桡侧腕屈肌 (FCR) 积分肌电 IEMG",
    "fds_iemg": "指浅屈肌 (FDS) 积分肌电 IEMG",
    "ecu_iemg": "尺侧腕伸肌 (ECU) 积分肌电 IEMG",
    "extensor_digitorum_iemg": "指伸肌 (Extensor Digitorum) 积分肌电 IEMG",
    "flexor_extensor_iemg_ratio": "屈伸肌 IEMG 比 = Σ屈肌 / Σ伸肌",
    "emg_burst_duration": "肌电爆发平均持续时间（秒）",
    "fcr_mdf": "桡侧腕屈肌 (FCR) 中位频率 MDF（疲劳代理）",
    "fds_mdf": "指浅屈肌 (FDS) 中位频率 MDF（疲劳代理）",
    "ecu_mdf": "尺侧腕伸肌 (ECU) 中位频率 MDF（疲劳代理）",
    "extensor_digitorum_mdf": "指伸肌中位频率 MDF（疲劳代理）",
    "movement_smoothness_sparc": "谱弧长运动平滑度（SPARC）",
    "range_of_motion_proxy": "4 传感器陀螺角速度范围均值 (p98-p2)",
    "tremor_index_3_6hz": "3–6 Hz 加速度相对功率（震颤指数）",
    "wrist_flexion_peak_velocity": "腕屈方向峰值角速度 (ECU 传感器主轴去偏置后负向 |p5|)",
    "wrist_extension_peak_velocity": "腕伸方向峰值角速度 (ECU 传感器主轴去偏置后正向 p95)",
    "finger_extension_peak_velocity": "伸指峰值角速度 (指伸肌处传感器陀螺幅值 p95)",
}

# 部位：优先取 reasonableness.SPECS，缺失项用此回退（仍取自 README 含义）
SITE_FALLBACK = {
    "prefrontal_theta_beta_ratio": "前额叶 Fp1/Fp2/Fz",
    "interhemispheric_motor_coherence": "健-患侧运动皮层 C3↔C4",
    "movement_mu_power_change": "运动皮层（EEG）",
    "movement_beta_power_change": "运动皮层（EEG）",
    "finger_co_contraction_index": "指浅屈肌 FDS vs 指伸肌",
    "fcr_iemg": "桡侧腕屈肌 FCR",
    "fds_iemg": "指浅屈肌 FDS（掌长肌电极位）",
    "ecu_iemg": "尺侧腕伸肌 ECU",
    "extensor_digitorum_iemg": "指伸肌",
    "flexor_extensor_iemg_ratio": "屈肌 vs 伸肌",
    "emg_burst_duration": "桡侧腕屈肌 FCR",
    "fcr_mdf": "桡侧腕屈肌 FCR",
    "fds_mdf": "指浅屈肌 FDS（掌长肌电极位）",
    "ecu_mdf": "尺侧腕伸肌 ECU",
    "extensor_digitorum_mdf": "指伸肌",
    "wrist_flexion_peak_velocity": "腕部 IMU 陀螺仪 (ECU)",
    "wrist_extension_peak_velocity": "腕部 IMU 陀螺仪 (ECU)",
    "finger_extension_peak_velocity": "指伸肌处 IMU 陀螺仪",
}

DIR_ZH = {"increase": "随康复↑（升高）", "decrease": "随康复↓（降低）", "n/a": "—"}
REFTYPE_ZH = {
    "healthy_norm": "健康常模",
    "directional_trend": "方向性趋势",
    "none": "无标准范围",
}

# 数据局限（取自 README“注意事项”小节，逐条）
LIMITATIONS = [
    "必须用 raw 信号：load_trial_raw 保留 EMG/IMU 绝对幅值（肌张力/震颤/IEMG 所需）；"
    "鲁棒 z-score 会抹掉它。",
    "EMG 4 通道：FCR / 指浅屈肌 (FDS) / ECU / 指伸肌。指浅屈肌 (FDS) 信号取自掌长肌 "
    "(Palmaris Longus) 电极位，按临床约定命名/使用为 FDS；据此腕屈/伸 (FCR vs ECU) 与"
    "指屈/伸 (FDS vs 指伸肌) 共收缩指数分别计算。",
    "运动相关 μ/β 功率变化：基于 EMG 包络划分高/低活动窗（非事件触发标记）；"
    "EEG 与 EMG 时钟独立、未跨模态精同步，其与 CMC、半球间相干的绝对值仅供队列内排名参考。",
    "腕屈/伸方向角速度 / 伸指速度：基于陀螺仪角速度（非关节角度测量），所有动作均为主动完成；"
    "腕屈/伸方向取 ECU 处传感器方差最大轴、高通去安装方向偏置后的带符号角速度正/负向峰值。",
    "中位频率 MDF：取各肌 Welch 功率谱 20–450 Hz 限带累计 50% 功率的频率，反映肌肉疲劳趋势；"
    "受电极位置/皮下脂肪影响，绝对值仅供队列内方向比较。",
    "mne 缺失时降级：EMG/IMU 标志物仍可计算，EEG 相关项为 NaN（聚合时 n_valid 计为 0）。",
    "真实 vs 模拟不可在绝对尺度上混比：模拟受试者 S6–S15 的 IMU 陀螺幅值比真实数据小约 2–3 个"
    "数量级，故 range_of_motion_proxy / wrist_flexion_peak_velocity / wrist_extension_peak_velocity"
    " / finger_extension_peak_velocity 等 IMU 绝对量在真实与模拟之间不可直接比较，宜各自组内比较。"
    "EEG（经同链处理）与 EMG IEMG 量级则大体可比；震颤、CCI、MDF 等相对/频率量跨组可作方向比较。",
]


# ----------------------------------------------------------------------------
# 数据装配
# ----------------------------------------------------------------------------
def load_sources():
    with open(REF_JSON, "r", encoding="utf-8") as f:
        ref = json.load(f)
    cohort = pd.read_csv(COHORT_CSV)
    return ref, cohort


def _healthy_str(hr):
    """把 healthy_reference 压成一段可读字符串。"""
    if not hr:
        return ""
    parts = []
    if hr.get("mean") is not None:
        parts.append(f"mean={hr['mean']}")
    if hr.get("sd") is not None:
        parts.append(f"sd={hr['sd']}")
    if hr.get("range") is not None:
        parts.append(f"range={hr['range'][0]}~{hr['range'][1]}")
    return ", ".join(parts)


def build_rows(ref):
    """为每项标志物逐项装配一行 dict（清单 + 参考范围共用）。"""
    bm = ref["biomarkers"]
    rows = []
    for name in BIOMARKER_NAMES:
        info = bm[name]
        spec = SPEC_BY_NAME.get(name)
        site = spec.site if spec else SITE_FALLBACK.get(name, "—")
        hr = info.get("healthy_reference")
        rows.append({
            "名称": name,
            "模态": MODALITY[name],
            "中文含义": MEANING_ZH[name],
            "单位": info.get("units", ""),
            "部位": site,
            "参考类型": REFTYPE_ZH.get(info.get("reference_type"), info.get("reference_type")),
            "随康复预期方向": DIR_ZH.get(
                info.get("expected_direction_with_recovery"),
                info.get("expected_direction_with_recovery"),
            ),
            "置信度": info.get("confidence", ""),
            # 参考范围 sheet 额外列
            "healthy_mean": (hr or {}).get("mean"),
            "healthy_sd": (hr or {}).get("sd"),
            "healthy_range": (hr or {}).get("range"),
            "healthy_str": _healthy_str(hr),
            "source": ", ".join(info.get("source", [])),
            "note_zh": info.get("note_zh", ""),
        })
    return rows


# ----------------------------------------------------------------------------
# Word
# ----------------------------------------------------------------------------
def make_docx(ref, cohort, rows):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # 全局中文字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    def add_table(headers, data_rows, col_keys):
        t = doc.add_table(rows=1, cols=len(headers))
        try:
            t.style = "Light Grid Accent 1"
        except Exception:
            t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for row in data_rows:
            cells = t.add_row().cells
            for i, k in enumerate(col_keys):
                cells[i].text = "" if row.get(k) is None else str(row.get(k))
        return t

    # ---- 标题页 ----
    title = doc.add_heading("三模态生物标志物提取模块 — 汇报说明", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("从 EEG / EMG / IMU 三模态信号提取临床可解释生物标志物，"
                            "并给出各标志物的文献参考范围")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(f"模块路径：analysis/02_biomarkers/    生成日期：{date.today().isoformat()}    "
                             f"生物标志物：{len(BIOMARKER_NAMES)} 项    队列：S1–S15")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 1 模块概述 ----
    doc.add_heading("一、模块概述", level=1)
    doc.add_paragraph(
        "本模块从单个病例的三模态信号（EEG 皮层 / EMG 外周肌肉 / IMU 运动学）提取一组"
        "有生理意义、可命名的生物标志物，并随病例临床标签（FMA-UE / BI / hand_tone(MAS) / "
        "hand_function(Brunnstrom)）一并输出。产出的是可解释的信号量，不做统计验证、更非临床诊断。"
    )
    for line in [
        "适用范围：支持 S1–S15。真实受试者仅 5 名（S1–S5，EEG 为 .bdf）；S6–S15 为模拟数据"
        "（EEG 为 .csv，走同链但不 z-score 的等效预处理，保证脑电标志物可比）。",
        "临床标签：统一取自 patient_rehab_suggestions_15subjects.json 的 labels（覆盖 S1–S15）。",
        "trial 定义：每个 trial = 一个主动抓握类动作的连续记录（抓握 / 勾状抓握 / 捏笔 / 拇指内收 / "
        "握圆筒 / 握球；S1 仅 4 试次）。所有动作均为主动完成。",
        "患侧映射：受损半球 = 患手对侧（患手 R → 左皮层 C3 簇），代码按 affected_side 字段双向通用。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ---- 2 标志物总表（分模态）----
    doc.add_heading(f"二、生物标志物清单（共 {len(BIOMARKER_NAMES)} 项）", level=1)
    headers = ["名称", "中文含义", "单位", "部位", "参考类型", "随康复预期方向", "置信度"]
    keys = headers
    for modality, zh in [("EEG", "EEG（皮层）"), ("EMG", "EMG（外周 / 肌肉）"), ("IMU", "IMU（运动学）")]:
        sub_rows = [r for r in rows if r["模态"] == modality]
        doc.add_heading(f"{zh} — {len(sub_rows)} 项", level=2)
        add_table(headers, sub_rows, keys)

    # ---- 3 参考范围与文献依据 ----
    doc.add_heading("三、参考范围与文献依据", level=1)
    doc.add_paragraph(
        "参考范围严格取自公开文献，分层优先 Brunnstrom，无分层依据时给健康成人常模。"
        "每项标注 reference_type：")
    for k, desc in [
        ("健康常模 (healthy_norm)", "有健康常模数值，如 SPARC、共收缩指数 CCI。"),
        ("方向性趋势 (directional_trend)", "仅方向性，如 CMC、前额叶 θ/β、PAI、震颤。"),
        ("无标准范围 (none)", "设备/协议特异量，文献无标准范围——多数 EMG 绝对量、IEMG、IMU 陀螺量、运动相关功率变化。"),
    ]:
        doc.add_paragraph(f"{k}：{desc}", style="List Bullet")

    doc.add_paragraph(
        "⚠️ 本模块若干量纲（原始 EMG 电压、IMU 陀螺 deg/s、加速度 SPARC）与文献测量方式不同，"
        "文献值即使存在也不可在绝对尺度直接套用，仅供方向参考。绝非临床诊断阈值。")

    doc.add_heading("有文献依据的标志物", level=2)
    cited = [r for r in rows if r["source"]]
    cite_rows = [{
        "名称": r["名称"],
        "健康参考": r["healthy_str"] or "（仅方向）",
        "随康复预期方向": r["随康复预期方向"],
        "文献": r["source"],
        "说明": r["note_zh"],
    } for r in cited]
    add_table(["名称", "健康参考", "随康复预期方向", "文献", "说明"],
              cite_rows, ["名称", "健康参考", "随康复预期方向", "文献", "说明"])

    doc.add_heading("参考文献", level=2)
    for i, refitem in enumerate(ref["references"], 1):
        doc.add_paragraph(
            f"[{refitem['id']}] {refitem['citation']}  {refitem.get('url', '')}",
            style="List Number")

    # ---- 4 合理性核对 ----
    doc.add_heading("四、合理性核对（reasonableness）", level=1)
    doc.add_paragraph(
        "对部分标志物，模块会把其数值与该病例临床标签做方向一致性核对：每个标志物对某个标签有"
        "预期方向（如 PAI 高 → 预期 FMA 低）。在 S1–S5 真实队列内，比较该受试者「标志物百分位」"
        "与「由标签推出的预期百分位」是否吻合，给出 合理 / 部分合理 / 不合理。")
    note = doc.add_paragraph("免责声明：" + DISCLAIMER)
    for r in note.runs:
        r.italic = True

    # ---- 5 数据局限与免责声明 ----
    doc.add_heading("五、数据局限与免责声明", level=1)
    for line in LIMITATIONS:
        doc.add_paragraph(line, style="List Bullet")
    disc = doc.add_paragraph(
        "总免责声明：真实受试者仅 5 名（S1–S5），其余 S6–S15 为模拟数据。本工具产出的是可解释的"
        "信号量，不做统计验证、更非临床诊断；参考范围绝非临床诊断阈值。")
    for r in disc.runs:
        r.bold = True

    os.makedirs(REPORT_DIR, exist_ok=True)
    doc.save(DOCX_PATH)
    return DOCX_PATH


# ----------------------------------------------------------------------------
# Excel
# ----------------------------------------------------------------------------
def make_xlsx(ref, cohort, rows):
    from openpyxl.styles import Font, Alignment
    from openpyxl.utils import get_column_letter

    os.makedirs(REPORT_DIR, exist_ok=True)

    # sheet 1：标志物清单
    df_list = pd.DataFrame([{
        "名称": r["名称"], "模态": r["模态"], "中文含义": r["中文含义"], "单位": r["单位"],
        "部位": r["部位"], "参考类型": r["参考类型"], "随康复预期方向": r["随康复预期方向"],
        "置信度": r["置信度"],
    } for r in rows])

    # sheet 2：参考范围
    df_ref = pd.DataFrame([{
        "名称": r["名称"], "单位": r["单位"], "参考类型": r["参考类型"],
        "healthy_mean": r["healthy_mean"], "healthy_sd": r["healthy_sd"],
        "healthy_range": "" if r["healthy_range"] is None else str(r["healthy_range"]),
        "随康复预期方向": r["随康复预期方向"], "置信度": r["置信度"],
        "文献": r["source"], "说明 note_zh": r["note_zh"],
    } for r in rows])

    # sheet 3：队列数值（直接来自 CSV）
    df_cohort = cohort.copy()

    # sheet 4：参考文献
    df_cite = pd.DataFrame([{
        "id": x["id"], "citation": x["citation"], "url": x.get("url", ""),
    } for x in ref["references"]])

    # sheet 5：数据局限说明
    limit_rows = [{"条目": f"局限 {i}", "说明": t} for i, t in enumerate(LIMITATIONS, 1)]
    limit_rows.append({"条目": "scale_caveat", "说明": ref["meta"].get("scale_caveat", "")})
    limit_rows.append({"条目": "disclaimer", "说明": ref["meta"].get("disclaimer", "")})
    limit_rows.append({"条目": "reasonableness 免责", "说明": DISCLAIMER})
    df_limit = pd.DataFrame(limit_rows)

    sheets = {
        "标志物清单": df_list,
        "参考范围": df_ref,
        "队列数值 S1-S15": df_cohort,
        "参考文献": df_cite,
        "数据局限说明": df_limit,
    }

    with pd.ExcelWriter(XLSX_PATH, engine="openpyxl") as writer:
        for sheet_name, df in sheets.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)

        # 美化：首行加粗、冻结首行、列宽自适应
        wb = writer.book
        for sheet_name, df in sheets.items():
            ws = wb[sheet_name]
            ws.freeze_panes = "A2"
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(vertical="center")
            for col_idx, col in enumerate(df.columns, 1):
                max_len = max(
                    [len(str(col))] + [len(str(v)) for v in df[col].astype(str).tolist()]
                )
                # 中文按 ~1.6 倍宽，封顶 60
                width = min(60, int(max_len * 1.6) + 2)
                ws.column_dimensions[get_column_letter(col_idx)].width = width

    return XLSX_PATH


# ----------------------------------------------------------------------------
def main():
    ref, cohort = load_sources()
    rows = build_rows(ref)
    assert len(rows) == len(BIOMARKER_NAMES), "标志物条数应与 BIOMARKER_NAMES 一致"

    docx_path = make_docx(ref, cohort, rows)
    xlsx_path = make_xlsx(ref, cohort, rows)

    print("✅ 已生成汇报材料：")
    print(f"   Word :  {docx_path}")
    print(f"   Excel:  {xlsx_path}")
    print(f"   标志物 {len(rows)} 项｜文献 {len(ref['references'])} 条｜队列 {len(cohort)} 人")


if __name__ == "__main__":
    main()
